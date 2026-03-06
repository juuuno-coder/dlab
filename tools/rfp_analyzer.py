"""
공고문/제안요청서 분석기
HWP, PDF, DOCX 파일에서 핵심 정보 추출
"""

import sys
import json
import re
from pathlib import Path
from typing import Dict, List, Any

# PDF 파싱
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# DOCX 파싱
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# HWP 파싱 (olefile 사용)
try:
    import olefile
    HAS_HWP = True
except ImportError:
    HAS_HWP = False


class RFPAnalyzer:
    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.file_ext = self.file_path.suffix.lower()
        self.raw_text = ""
        self.analysis = {
            "메타정보": {},
            "사업개요": {},
            "예산": {},
            "일정": {},
            "평가기준": {},
            "핵심요구사항": [],
            "제출서류": [],
            "특이사항": []
        }

    def extract_text(self):
        """파일 형식에 따라 텍스트 추출"""
        if self.file_ext == '.pdf':
            return self._extract_from_pdf()
        elif self.file_ext == '.docx':
            return self._extract_from_docx()
        elif self.file_ext == '.hwp':
            return self._extract_from_hwp()
        else:
            raise ValueError(f"지원하지 않는 파일 형식: {self.file_ext}")

    def _extract_from_pdf(self):
        """PDF에서 텍스트 추출"""
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF가 설치되지 않았습니다: pip install PyMuPDF")

        doc = fitz.open(str(self.file_path))
        text_parts = []

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            text_parts.append(text)

        self.raw_text = '\n'.join(text_parts)
        return self.raw_text

    def _extract_from_docx(self):
        """DOCX에서 텍스트 추출"""
        if not HAS_DOCX:
            raise ImportError("python-docx가 설치되지 않았습니다: pip install python-docx")

        doc = Document(str(self.file_path))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 표 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join([cell.text for cell in row.cells])
                text_parts.append(row_text)

        self.raw_text = '\n'.join(text_parts)
        return self.raw_text

    def _extract_from_hwp(self):
        """HWP에서 텍스트 추출 (간단한 방식)"""
        if not HAS_HWP:
            raise ImportError("olefile이 설치되지 않았습니다: pip install olefile")

        # HWP는 복합 문서 형식이므로 완벽한 파싱은 어려움
        # 기본적인 텍스트만 추출
        try:
            ole = olefile.OleFileIO(str(self.file_path))

            # BodyText 섹션에서 텍스트 추출 시도
            text_parts = []
            for entry in ole.listdir():
                if 'BodyText' in entry:
                    stream = ole.openstream(entry)
                    data = stream.read()
                    # 한글 인코딩 시도 (UTF-16LE)
                    try:
                        text = data.decode('utf-16le', errors='ignore')
                        text_parts.append(text)
                    except:
                        pass

            self.raw_text = '\n'.join(text_parts)

            # 텍스트가 없으면 경고
            if not self.raw_text.strip():
                print("⚠️ HWP 파일에서 텍스트를 추출하지 못했습니다.")
                print("   HWP 파일을 PDF로 변환한 후 다시 시도해주세요.")

            return self.raw_text

        except Exception as e:
            raise RuntimeError(f"HWP 파일 파싱 실패: {e}")

    def analyze(self):
        """텍스트 분석하여 핵심 정보 추출"""
        if not self.raw_text:
            self.extract_text()

        # 1. 사업명 추출
        self._extract_project_name()

        # 2. 발주기관 추출
        self._extract_client()

        # 3. 예산 추출
        self._extract_budget()

        # 4. 일정 추출
        self._extract_schedule()

        # 5. 평가기준 추출
        self._extract_evaluation_criteria()

        # 6. 핵심 요구사항 추출
        self._extract_requirements()

        # 7. 제출서류 추출
        self._extract_documents()

        return self.analysis

    def _extract_project_name(self):
        """사업명 추출"""
        patterns = [
            r'사업명\s*[:：]\s*(.+)',
            r'과제명\s*[:：]\s*(.+)',
            r'용역명\s*[:：]\s*(.+)',
            r'^\s*(.+?(?:사업|용역|과제|행사|전시회|축제))\s*$'
        ]

        for pattern in patterns:
            match = re.search(pattern, self.raw_text, re.MULTILINE)
            if match:
                name = match.group(1).strip()
                # 너무 긴 경우 첫 줄만
                if '\n' in name:
                    name = name.split('\n')[0].strip()
                self.analysis["사업개요"]["사업명"] = name
                break

    def _extract_client(self):
        """발주기관 추출"""
        patterns = [
            r'발주기관\s*[:：]\s*(.+)',
            r'발주처\s*[:：]\s*(.+)',
            r'주최\s*[:：]\s*(.+)',
            r'주관\s*[:：]\s*(.+)'
        ]

        for pattern in patterns:
            match = re.search(pattern, self.raw_text)
            if match:
                self.analysis["사업개요"]["발주기관"] = match.group(1).strip()
                break

    def _extract_budget(self):
        """예산 추출"""
        # 금액 패턴: 숫자 + 억원, 천만원, 만원 등
        patterns = [
            r'예산\s*[:：]?\s*([0-9,]+(?:\s*억|\s*천만|\s*만)?원?)',
            r'총\s*사업비\s*[:：]?\s*([0-9,]+(?:\s*억|\s*천만|\s*만)?원?)',
            r'계약금액\s*[:：]?\s*([0-9,]+(?:\s*억|\s*천만|\s*만)?원?)',
        ]

        for pattern in patterns:
            match = re.search(pattern, self.raw_text)
            if match:
                budget_str = match.group(1).strip()
                self.analysis["예산"]["총예산"] = budget_str

                # 숫자 파싱 시도
                try:
                    # 억원 단위 변환
                    if '억' in budget_str:
                        num = re.search(r'([0-9,]+)', budget_str).group(1)
                        num = int(num.replace(',', ''))
                        self.analysis["예산"]["금액(원)"] = num * 100000000
                    elif '만' in budget_str:
                        num = re.search(r'([0-9,]+)', budget_str).group(1)
                        num = int(num.replace(',', ''))
                        self.analysis["예산"]["금액(원)"] = num * 10000
                except:
                    pass

                break

    def _extract_schedule(self):
        """일정 추출"""
        # 날짜 패턴: YYYY.MM.DD, YYYY-MM-DD, YYYY/MM/DD
        date_pattern = r'(\d{4}[.-/]\d{1,2}[.-/]\d{1,2})'

        # 사업기간 패턴
        period_patterns = [
            r'사업기간\s*[:：]?\s*' + date_pattern + r'\s*[~∼~-]\s*' + date_pattern,
            r'용역기간\s*[:：]?\s*' + date_pattern + r'\s*[~∼~-]\s*' + date_pattern,
            r'계약기간\s*[:：]?\s*' + date_pattern + r'\s*[~∼~-]\s*' + date_pattern,
        ]

        for pattern in period_patterns:
            match = re.search(pattern, self.raw_text)
            if match:
                self.analysis["일정"]["시작일"] = match.group(1)
                self.analysis["일정"]["종료일"] = match.group(2)
                self.analysis["일정"]["기간"] = f"{match.group(1)} ~ {match.group(2)}"
                break

    def _extract_evaluation_criteria(self):
        """평가기준 추출"""
        # 평가 항목: 숫자 + 점
        eval_pattern = r'(.+?)\s*[:：]?\s*(\d+)\s*점'

        matches = re.findall(eval_pattern, self.raw_text)

        criteria = {}
        for item, score in matches:
            item = item.strip()
            # 너무 긴 항목명은 제외
            if len(item) < 50 and len(item) > 2:
                # 평가 관련 키워드가 있는 경우만
                if any(kw in item for kw in ['평가', '능력', '계획', '수행', '실적', '제안']):
                    criteria[item] = int(score)

        self.analysis["평가기준"] = criteria

    def _extract_requirements(self):
        """핵심 요구사항 추출"""
        # 요구사항 키워드 섹션 찾기
        requirement_sections = [
            r'(?:과업|사업)\s*내용',
            r'주요\s*(?:내용|업무)',
            r'추진\s*내용',
            r'세부\s*과업'
        ]

        requirements = []

        for section_pattern in requirement_sections:
            # 섹션 시작 위치 찾기
            match = re.search(section_pattern, self.raw_text)
            if match:
                start = match.end()
                # 섹션 종료까지 (다음 대제목 또는 500자)
                section_text = self.raw_text[start:start+1000]

                # 번호 리스트 추출 (1. 2. 3. 또는 가. 나. 다.)
                list_items = re.findall(r'(?:^|\n)\s*(?:\d+\.|[가-힣]\.)\s*(.+)', section_text)

                for item in list_items[:10]:  # 최대 10개
                    item = item.strip()
                    if len(item) > 5 and len(item) < 200:
                        requirements.append(item)

        self.analysis["핵심요구사항"] = requirements

    def _extract_documents(self):
        """제출서류 추출"""
        # 제출서류 섹션 찾기
        doc_section_pattern = r'제출\s*(?:서류|자료)'

        match = re.search(doc_section_pattern, self.raw_text)
        if match:
            start = match.end()
            section_text = self.raw_text[start:start+500]

            # 번호 리스트 추출
            list_items = re.findall(r'(?:^|\n)\s*(?:\d+\.|[가-힣]\.)\s*(.+)', section_text)

            documents = []
            for item in list_items[:15]:  # 최대 15개
                item = item.strip()
                if len(item) > 3 and len(item) < 100:
                    documents.append(item)

            self.analysis["제출서류"] = documents

    def save_json(self, output_path):
        """분석 결과를 JSON으로 저장"""
        output_file = Path(output_path)

        # 메타정보 추가
        self.analysis["메타정보"] = {
            "원본파일": self.file_path.name,
            "파일형식": self.file_ext,
            "분석일시": str(Path.ctime(self.file_path))
        }

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(self.analysis, f, ensure_ascii=False, indent=2)

        print(f"\n✅ 분석 결과 저장: {output_file}")
        return output_file

    def print_summary(self):
        """분석 결과 요약 출력"""
        print("\n" + "=" * 80)
        print("📋 공고문 분석 결과")
        print("=" * 80)

        if self.analysis["사업개요"].get("사업명"):
            print(f"\n📌 사업명: {self.analysis['사업개요']['사업명']}")

        if self.analysis["사업개요"].get("발주기관"):
            print(f"🏢 발주기관: {self.analysis['사업개요']['발주기관']}")

        if self.analysis["예산"].get("총예산"):
            print(f"💰 예산: {self.analysis['예산']['총예산']}")

        if self.analysis["일정"].get("기간"):
            print(f"📅 기간: {self.analysis['일정']['기간']}")

        if self.analysis["평가기준"]:
            print(f"\n📊 평가기준:")
            for item, score in self.analysis["평가기준"].items():
                print(f"  - {item}: {score}점")

        if self.analysis["핵심요구사항"]:
            print(f"\n🎯 핵심 요구사항 ({len(self.analysis['핵심요구사항'])}개):")
            for i, req in enumerate(self.analysis["핵심요구사항"][:5], 1):
                print(f"  {i}. {req}")
            if len(self.analysis["핵심요구사항"]) > 5:
                print(f"  ... 외 {len(self.analysis['핵심요구사항']) - 5}개")

        if self.analysis["제출서류"]:
            print(f"\n📄 제출서류 ({len(self.analysis['제출서류'])}개):")
            for i, doc in enumerate(self.analysis["제출서류"][:5], 1):
                print(f"  {i}. {doc}")
            if len(self.analysis["제출서류"]) > 5:
                print(f"  ... 외 {len(self.analysis['제출서류']) - 5}개")

        print("\n" + "=" * 80)


def main():
    if len(sys.argv) < 2:
        print("사용법: python rfp_analyzer.py <공고문파일> [출력JSON경로]")
        print("예시: python rfp_analyzer.py 입찰공고.pdf rfp_analysis.json")
        print("\n지원 형식: PDF, DOCX, HWP")
        sys.exit(1)

    file_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else 'rfp_analysis.json'

    try:
        analyzer = RFPAnalyzer(file_path)

        print(f"\n📖 파일 읽기 중: {file_path}")
        analyzer.extract_text()

        print(f"📝 텍스트 추출 완료 ({len(analyzer.raw_text)} 문자)")

        print(f"🔍 분석 중...")
        analyzer.analyze()

        analyzer.print_summary()
        analyzer.save_json(output_path)

    except Exception as e:
        print(f"\n❌ 오류 발생: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
